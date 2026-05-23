from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# Helper functions
def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bold_then_normal(bold_text, normal_text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(normal_text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# ============ HEADER ============
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run('LEMBAR JAWABAN ASSESSMENT (ATS)')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('TI41254 Software Engineering 2\nSemester Genap T.A. 2025-2026')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

doc.add_paragraph()  # spacer

# Identity table
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
cells = [
    ('Kelompok', '7 — API Gateway / Integrator'),
    ('Anggota', 'Richard Firmansyah (714240047)\nZidan Hairra Ramadhan (714240061)'),
    ('Kelas', 'TI-126L'),
    ('Tech Stack', 'Node.js, Express v5, EJS, JWT (jsonwebtoken), Axios'),
]
for i, (k, v) in enumerate(cells):
    c0 = table.cell(i, 0)
    c1 = table.cell(i, 1)
    c0.text = k
    c1.text = v
    for c in [c0, c1]:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
    # Bold the key column
    for run in c0.paragraphs[0].runs:
        run.bold = True

doc.add_paragraph()

# ============ SOAL 1 ============
add_heading_custom('SOAL 1 — Nama Aplikasi dan Deskripsi', level=1)

add_para('Aplikasi yang kami kembangkan bernama API Gateway / Integrator. Sesuai dengan instruksi ekosistem UMKM Digital ULBI, aplikasi kami bertindak sebagai middleware orchestrator atau pintu masuk tunggal (single entry point) untuk seluruh komunikasi antar 6 aplikasi kelompok lainnya.')

add_para('Secara konsep, aplikasi kami tidak memiliki logika bisnis utama layaknya marketplace atau sistem gudang. Peran utama kami murni sebagai lapisan infrastruktur yang mengatur lalu lintas request, memastikan keamanan lewat validasi token JWT, mencatat semua riwayat transaksi, serta memotong fee gateway sebesar 0.5% secara otomatis untuk setiap transaksi komersial.')

add_para('Secara teknis, kami membangun 4 fitur utama sesuai spesifikasi Doc1–Doc6: Routing API (/integrator/routing_api) untuk meneruskan request ke 6 service, Validasi Request (/integrator/validasi_request) untuk otentikasi JWT, Logging (/integrator/logging) untuk mencatat setiap aktivitas, dan Biaya Layanan Integrasi (/integrator/biaya_layanan_integrasi) untuk mengelola fee 0.5% per transaksi. Seluruh konfigurasi service URL disimpan di file .env agar mudah diubah tanpa menyentuh kode.')

# ============ SOAL 2 ============
add_heading_custom('SOAL 2 — Proses Transaksi End-to-End [Bobot 50]', level=1)

add_para('Sebagai pengembang API Gateway, kami menyadari bahwa aplikasi kami adalah perantara krusial. Mari ambil contoh kasus utama: Marketplace melakukan checkout yang membutuhkan pemotongan saldo di SmartBank. Seluruh proses ini wajib melewati gateway kami sesuai Doc4 Aturan No.5.')

# 2.1
add_heading_custom('1) Input Utama yang Diterima', level=2)
add_para('Ketika ada request masuk, misalnya Marketplace memanggil POST /integrator/smartbank/pembayaran_transaksi, hal pertama yang kami tangkap adalah Header Authorization berisi token JWT, parameter :service dari URL (misalnya "smartbank"), wildcard path setelahnya ("pembayaran_transaksi"), serta body request yang berisi data seperti user_id, amount, dan parameter transaksi. Middleware logger.js kami juga secara otomatis mencatat IP pengirim, HTTP method, dan URL tujuan lengkap.')

# 2.2
add_heading_custom('2) API yang Dipanggil ke Sistem Lain', level=2)
add_para('Setelah lolos pengecekan token, sistem kami melakukan dua pemanggilan API secara berurutan. Pertama, sistem memanggil API ke SmartBank secara otomatis untuk memotong fee 0.5% dari amount transaksi melalui POST ke endpoint /smartbank/pembayaran_transaksi. Kami mengirimkan data berupa user_id, jumlah fee yang dipotong, parameter keterangan "Biaya Layanan Integrasi", serta original_amount agar SmartBank tahu ini potongan fee, bukan transaksi utama. Jika pemotongan fee ini berhasil, barulah sistem melakukan pemanggilan API kedua menggunakan Axios untuk meneruskan request asli beserta datanya ke service tujuan yang sebenarnya.')

# 2.3
add_heading_custom('3) Data yang Dikirim dan Diterima', level=2)
add_para('Untuk fee, data yang kami kirim ke SmartBank berisi: user_id pengguna, amount senilai fee (misal Rp 250 dari transaksi Rp 50.000), parameter berisi "Biaya Layanan Integrasi (Gateway Fee 0.5%)", source bertuliskan "integrator", dan original_amount yaitu Rp 50.000. Untuk request yang di-forward, kami meneruskan seluruh body asli beserta header Authorization dan Content-Type.')

add_para('Response yang kami kembalikan ke aplikasi asal selalu kami bungkus dengan objek integrator_info yang berisi informasi service tujuan, persentase fee, jumlah fee yang terpotong, status fee (terpotong/gagal_potong), dan URL service yang dituju. Di dalamnya juga terdapat data respons asli dari service tujuan. Dengan cara ini, aplikasi pengirim tahu persis apa yang terjadi di gateway kami.')

# 2.4
add_heading_custom('4) Mekanisme Validasi JWT/Token', level=2)
add_para('Untuk menjaga keamanan ekosistem, kami memusatkan fungsi validasi di file terpisah bernama middleware/auth.js. Prosesnya ada 5 langkah: pertama, sistem mengekstrak token dari header Authorization: Bearer <token>. Kedua, jika tidak ada token, kami langsung tolak dengan HTTP 401 beserta pesan bahwa token diperlukan. Ketiga, kami memverifikasi token menggunakan jwt.verify() dengan secret key yang tersimpan di file .env. Keempat, jika valid, payload token di-decode dan disimpan ke req.user — berisi user_id, name, npm, role, dan waktu pembuatan token. User_id ini juga langsung kami perbarui di log entry untuk keperluan audit. Kelima, jika token sudah kedaluwarsa atau tidak valid, kami kembalikan HTTP 403 beserta detail error-nya.')

add_para('Satu hal penting dalam arsitektur kami: auth middleware ini hanya dipasang pada route /integrator/* melalui app.use(\'/integrator\', loggerMiddleware, validateRequest, gatewayRoutes). Endpoint publik seperti /api/status dan /generate-test-token sengaja tidak memerlukan auth agar bisa diakses untuk monitoring dan pembuatan token awal.')

# 2.5
add_heading_custom('5) Risiko Inkonsistensi Data', level=2)
add_para('Tentu saja ada risiko inkonsistensi data yang kami antisipasi. Risiko terbesarnya adalah jika fee 0.5% sudah berhasil dipotong ke SmartBank, namun saat sistem kami mencoba mem-forward request ke service tujuan, layanannya sedang down atau timeout. Jika ini terjadi, user akan kehilangan fee tanpa mendapatkan layanannya — dan saat ini kami belum memiliki mekanisme rollback otomatis.')

add_para('Risiko kedua adalah double transaction: jika client melakukan retry karena timeout, fee bisa dipotong dua kali karena kami belum memiliki idempotency check. Risiko ketiga terkait penyimpanan — log kami tersimpan di global.requestLogs yang bersifat in-memory, sehingga jika server restart, seluruh data log dan revenue record hilang. Terakhir, ada potensi race condition ketika multiple request masuk bersamaan dan mengakses requestLogs[length-1] yang bisa menyebabkan log entry tertukar antar request.')

# 2.6
add_heading_custom('6) Dampak Jika Salah Satu Aplikasi Lain Gagal', level=2)
add_para('Dampaknya berbeda tergantung aplikasi mana yang gagal. Jika SmartBank down, fee gateway kami gagal dipotong (tercatat sebagai feeStatus = \'gagal_potong\'), namun kami memilih untuk tetap meneruskan request ke service tujuan karena try-catch kami memisahkan proses fee dan forward — filosofinya, lebih baik transaksi utama jalan meski fee belum terpotong.')

add_para('Jika Marketplace, POS, SupplierHub, atau LogistiKita yang down, gateway kami akan mengembalikan HTTP 502 Bad Gateway dengan pesan error yang jelas. Log kami mencatat status ERROR beserta HTTP status code dari service yang gagal. Untuk UMKM Insight, dampaknya relatif minimal karena sifatnya read-only sesuai Doc4 Aturan 7 — kegagalannya tidak mempengaruhi transaksi ekonomi.')

add_para('Yang paling kritis: karena arsitektur kami bersifat terpusat, jika API Gateway sendiri yang mengalami down, maka seluruh komunikasi antar kelompok di ekosistem ini akan lumpuh total. Ini yang disebut Single Point of Failure dan menjadi kelemahan terbesar arsitektur gateway terpusat.')

# 2.7
add_heading_custom('7) Strategi agar Sistem Tetap Robust', level=2)
add_para('Untuk membuat sistem tetap robust, pendekatan utama kami adalah memisahkan blok try-catch antara pemotongan fee dan penerusan request di gateway.js. Jadi, jika error terjadi saat fee collection, request utama transaksi tidak langsung mati — dicatat dengan status gagal_potong namun operasi bisnis tetap bisa berjalan. Ini kami sebut sebagai fail-safe fee handling.')

add_para('Kami juga memberlakukan timeout 10 detik pada setiap request Axios agar server tidak menggantung menunggu balasan dari service yang tidak responsif. Selain itu, middleware chain kami (Logger → Auth → Gateway) dirancang independen — kegagalan di satu layer tidak merusak layer lainnya. Konfigurasi URL service menggunakan environment variable di .env, sehingga jika ada service yang pindah server, kami cukup ubah konfigurasi tanpa sentuh kode.')

add_para('Kami juga menyediakan endpoint Demo/Simulasi (/api/demo/simulate) yang selalu mengembalikan respons sukses meskipun service lain belum berjalan. Ini menjadi fallback penting saat presentasi atau testing.')

add_para('Strategi ideal yang belum kami implementasikan namun sangat kami sadari perlu ditambahkan: Circuit Breaker untuk memutus jalur ke service yang sudah gagal berkali-kali, Idempotency Key untuk mencegah double transaction, Saga Pattern untuk rollback fee jika forward gagal, dan Persistent Storage (database/Redis) untuk menggantikan in-memory log.')

# ============ SOAL 3 ============
add_heading_custom('SOAL 3 — Respons Terhadap Lonjakan Transaksi [Bobot 50]', level=1)

add_para('Menghadapi skenario traffic tinggi di mana SmartBank mengalami delay, Marketplace terus menerima pesanan, SupplierHub kekurangan stok, dan LogistiKita terlambat sinkronisasi ongkir, kami harus menerapkan arsitektur yang tangguh agar ekosistem tidak mengalami cascade failure. Analisis berikut kami lakukan dari perspektif API Gateway sebagai orchestrator.')

# 3.1
add_heading_custom('1) Menjaga Konsistensi Transaksi Ekonomi', level=2)
add_para('Ketika SmartBank delay, gateway kami juga ikut menunggu saat mencoba memotong fee. Untungnya, kami sudah memberlakukan timeout 10 detik, sehingga request tidak menggantung selamanya. Fee dan forward kami jalankan secara sekuensial (fee dulu, baru forward) sehingga urutan proses selalu benar.')

add_para('Untuk mengatasi kondisi delay berkelanjutan, solusi ideal kami adalah menerapkan request queuing berbasis Redis. Request yang masuk tidak langsung ditolak, melainkan disimpan di queue. Fee collection dilakukan oleh worker async yang terpisah sehingga tidak memblokir forward request ke service tujuan. Gateway menerima request, mencatat di log, dan mengembalikan status "sedang diproses" ke client — pendekatan eventual consistency ini menjaga alur ekonomi tetap berjalan meski ada delay.')

# 3.2
add_heading_custom('2) Mencegah Double Transaction', level=2)
add_para('Karena SmartBank mengalami delay, sangat mungkin pengguna Marketplace panik dan menekan tombol checkout berkali-kali. Untuk mencegah terpotongnya saldo atau fee lebih dari satu kali, solusi strategis kami adalah mengimplementasikan Idempotency Key. Dengan menanamkan middleware pengecekan di awal, jika ada request duplikat dengan kunci yang sama masuk di waktu berdekatan, Gateway cukup membalasnya dengan respons dari log yang tersimpan tanpa memproses ulang pemotongan fee atau mem-forward request ke service tujuan.')

add_para('Data log yang sudah kami miliki (id + timestamp + user_id) juga membantu deteksi duplikat dalam window waktu tertentu, sejalan dengan Doc6 Aturan No.15 tentang cooldown transaksi 10-30 detik dan Aturan No.16 tentang maksimal 10 transaksi per hari.')

# 3.3
add_heading_custom('3) Mencegah Pengurangan Stok Palsu', level=2)
add_para('Untuk mencegah pengurangan stok palsu di SupplierHub saat pembayaran di SmartBank delay, kami sangat merekomendasikan penerapan Saga Pattern yang diorkestrasi oleh Gateway. Alurnya: pertama, SupplierHub mengubah stok menjadi status RESERVED (bukan langsung dikurangi). Kedua, Gateway meneruskan pembayaran ke SmartBank. Jika SmartBank merespons pembayaran berhasil, baru SupplierHub mengonfirmasi stok menjadi CONFIRMED. Namun jika pembayaran gagal, Gateway mengirim compensating transaction ke SupplierHub untuk mengembalikan stok dari RESERVED kembali ke AVAILABLE.')

add_para('Selain itu, gateway juga bisa memvalidasi respons dari SupplierHub sebelum melanjutkan ke langkah berikutnya. Jika SupplierHub sudah mengembalikan error stok habis (HTTP 409), gateway langsung menghentikan alur transaksi tanpa memotong fee atau melanjutkan ke SmartBank.')

# 3.4
add_heading_custom('4) Menjaga Sistem Tetap Scalable', level=2)
add_para('Secara arsitektur, gateway kami sudah bersifat stateless — setiap request ditangani secara independen tanpa menyimpan session per user. Ini berarti gateway kami bisa langsung di-scale horizontal dengan menempatkan beberapa instance di belakang load balancer seperti Nginx atau HAProxy.')

add_para('Namun ada satu hambatan: global.requestLogs kami saat ini tersimpan di memori proses. Jika kami deploy 2 instance, masing-masing punya log terpisah dan dashboard menjadi tidak akurat. Solusinya adalah memindahkan penyimpanan log ke Redis atau database, sehingga semua instance membaca dan menulis ke tempat yang sama.')

add_para('Untuk mengendalikan beban saat lonjakan, kami juga perlu menerapkan rate limiting per user — membatasi sesuai Doc6 Aturan 15 (cooldown 10-30 detik) dan Aturan 16 (max 10 transaksi per hari). Dengan begitu, gateway tidak kewalahan memproses request yang seharusnya memang ditahan.')

# 3.5
add_heading_custom('5) Memberikan Feedback yang Jelas kepada User', level=2)
add_para('Saat ini gateway kami sudah mengembalikan respons terstruktur yang cukup informatif. Setiap response berisi status (success/error), message yang menjelaskan apa yang terjadi, integrator_info yang menunjukkan detail fee dan forwarding, serta error_detail jika terjadi kegagalan. Dashboard admin kami di /dashboard juga menampilkan statistik real-time: total revenue, total traffic, jumlah sukses, dan jumlah error.')

add_para('Untuk kondisi lonjakan, kami perlu menambahkan informasi tambahan dalam respons error, seperti estimated_retry_time agar client tahu kapan sebaiknya mencoba lagi, dan HTTP status code yang lebih spesifik — misalnya 504 Gateway Timeout jika SmartBank delay, 429 Too Many Requests jika user melebihi batas transaksi, dan 503 Service Unavailable jika gateway sedang overload.')

# 3.6
add_heading_custom('6) Mencegah Cascade Failure di Ekosistem', level=2)
add_para('Ini adalah risiko terbesar karena gateway adalah satu-satunya jalur komunikasi semua aplikasi. Jika SmartBank terus delay, antrean request di Axios gateway kami bisa menumpuk dan menghabiskan sumber daya server, membuat sistem kami ikut lumpuh — inilah cascade failure.')

add_para('Solusi pertama yang kami rekomendasikan adalah Circuit Breaker. Jika SmartBank sudah gagal merespons sebanyak 5 kali berturut-turut, gateway akan "membuka circuit" dan langsung mengembalikan error tanpa mencoba koneksi lagi selama 30 detik. Ini menghemat sumber daya yang seharusnya terbuang untuk menunggu service yang sedang bermasalah.')

add_para('Solusi kedua adalah Bulkhead Pattern — membatasi jumlah concurrent request ke tiap service. Misalnya SmartBank dibatasi maksimal 50 koneksi bersamaan. Jika sudah penuh, request selanjutnya langsung ditolak dengan HTTP 429. Dengan begitu, overload di SmartBank tidak menghabiskan seluruh kapasitas gateway yang seharusnya juga melayani Marketplace, POS, dan service lainnya.')

add_para('Yang sudah kami implementasikan adalah Graceful Degradation: jika fee collection gagal karena SmartBank error, gateway tetap meneruskan transaksi utama (feeStatus = \'gagal_potong\'). Ekosistem tidak berhenti hanya karena fee gagal dipotong. Ditambah timeout 10 detik yang mencegah thread tergantung selamanya.')

# -- Komponen Kritis --
add_heading_custom('Komponen Paling Kritis', level=2)
add_para('Dalam kondisi krisis seperti ini, komponen yang paling kritis dan harus kami prioritaskan adalah: pertama, JWT Auth Middleware (auth.js) karena ini adalah gerbang keamanan utama — tanpa ini, request tidak sah bisa masuk ke seluruh ekosistem. Kedua, Fee Calculator di gateway.js karena salah perhitungan akan merugikan user atau menghilangkan revenue gateway. Ketiga, Request Forwarder yang merupakan fungsi inti gateway — tanpa ini gateway tidak berguna. Keempat, Logger Middleware untuk audit trail, meskipun dalam kondisi darurat ini bisa diredamkan terlebih dahulu agar beban server berkurang.')

# -- Endpoint Prioritas --
add_heading_custom('Endpoint yang Harus Diprioritaskan', level=2)
add_para('Endpoint paling kritis adalah route forwarding utama (POST /integrator/:service/*) karena semua transaksi ekosistem melewatinya. Setelah itu, endpoint generate token karena tanpa token tidak ada yang bisa mengakses gateway. Kemudian /api/status sebagai health check untuk monitoring. Sementara endpoint pelaporan seperti /integrator/logging dan /api/logs bisa diturunkan prioritasnya saat lonjakan — lebih baik resource dialokasikan untuk menangani request transaksi.')

# -- Log Wajib --
add_heading_custom('Log yang Wajib Dicatat', level=2)
add_para('Berdasarkan implementasi logger.js kami, log wajib mencakup: waktu dan timestamp untuk kronologi audit, IP address untuk identifikasi sumber dan deteksi abuse, HTTP method dan URL tujuan untuk mengetahui operasi yang dilakukan, user_id untuk akuntabilitas per pengguna, service tujuan untuk tracking service mana yang dipanggil, status request (PENDING/SUCCESS/ERROR) untuk monitoring keberhasilan, fee_terpotong dan fee_status untuk audit keuangan, serta response HTTP status code untuk debugging.')

add_para('Khusus dalam kondisi lonjakan, kami perlu menambahkan pencatatan response time (latency per request) untuk mendeteksi service mana yang menyebabkan delay, serta circuit breaker state untuk mengetahui service mana yang sudah diputus jalurnya.')

# -- SOLID --
add_heading_custom('Penerapan Clean Architecture dan SOLID', level=2)

add_para('Kami merancang kode sedemikian rupa agar penerapan prinsip SOLID benar-benar terasa manfaatnya saat terjadi masalah lonjakan.')

add_bold_then_normal('Single Responsibility Principle (SRP): ', 'Kami memecah file secara spesifik — logger.js hanya fokus mencatat aktivitas, auth.js hanya mengurus JWT, dan gateway.js hanya menangani routing dan fee. Manfaatnya saat lonjakan: jika proses logging menyebabkan bottleneck saat traffic tinggi, kami bisa mematikan atau meredamkan fitur log saja tanpa merusak sistem otentikasi dan routing utama. Debugging juga lebih cepat karena kami tahu persis file mana yang bermasalah.')

add_bold_then_normal('Open/Closed Principle (OCP): ', 'Kami membuat pemetaan URL service secara dinamis menggunakan SERVICE_MAP yang membaca dari environment variable. Jika tim SmartBank memindahkan servernya untuk mengatasi delay, kami tidak perlu mengedit satu baris kode pun di gateway. Kami cukup mengubah URL di file .env, dan sistem langsung berjalan normal kembali tanpa proses restart yang berisiko. Menambah service baru ke ekosistem juga cukup tambah satu baris di .env — routing otomatis mendukungnya karena kami menggunakan dynamic route /:service/{*path}.')

add_bold_then_normal('Liskov Substitution Principle (LSP): ', 'Semua 6 service dalam SERVICE_MAP kami perlakukan secara identik oleh gateway. SmartBank, Marketplace, POS — semuanya melewati routing yang sama dan bisa saling digantikan URL-nya tanpa mengubah logika gateway. Jika SmartBank delay dan timnya menyediakan server cadangan, kami tinggal ganti URL-nya dan gateway tetap berfungsi normal.')

add_bold_then_normal('Interface Segregation Principle (ISP): ', 'Middleware kami bersifat komposabel dan tidak dipaksakan ke semua route. Endpoint /integrator/* melewati 3 middleware (logger, auth, gateway), sementara endpoint publik seperti /api/status dan /api/demo/simulate tidak memerlukan auth. Dengan pemisahan ini, saat lonjakan kami bisa rate-limit endpoint non-kritis tanpa mengganggu endpoint transaksi utama.')

add_bold_then_normal('Dependency Inversion Principle (DIP): ', 'Gateway kami bergantung pada abstraksi (URL di environment variable) — bukan pada implementasi konkret service tertentu. Jika kami perlu mengganti penyimpanan log dari in-memory ke Redis, kami cukup mengubah implementasi di logger.js tanpa menyentuh auth.js atau gateway.js, karena interface-nya tetap sama: terima request, catat log, panggil next().')

add_para('Penerapan Clean Architecture juga terlihat dari pemisahan layer: Presentation Layer (views/*.ejs untuk UI), Application Layer (server.js untuk konfigurasi route), Domain/Business Layer (routes/gateway.js untuk logika fee dan routing), dan Infrastructure Layer (middleware/*.js untuk JWT dan logging). Setiap layer independen, sehingga perubahan UI tidak mempengaruhi logika fee, dan sebaliknya.')

add_para('Manfaat nyata arsitektur ini saat lonjakan: ketika perlu menambah rate limiter, kami cukup buat satu file middleware baru dan mount di server.js — tidak ada satu baris pun di gateway.js atau auth.js yang perlu diubah. Risiko regression bug minimal karena perubahan terlokalisasi.')

# ============ FOOTER ============
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('— Akhir Lembar Jawaban —')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(11)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer2.add_run('API Gateway / Integrator — Kelompok 7\nRichard Firmansyah (714240047) & Zidan Hairra Ramadhan (714240061)\nTI41254 Software Engineering 2 — ULBI 2025-2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(10)
r.italic = True

# Save
output_path = r'c:\Users\achma\Documents\RPL_Integrator-main (1)\RPL_Integrator-main\ATS\Jawaban_ATS_Naratif_FINAL.docx'
doc.save(output_path)
print(f'File berhasil disimpan ke: {output_path}')
